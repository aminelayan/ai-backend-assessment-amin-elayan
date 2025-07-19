from docx import Document
from datetime import datetime
import os
import uuid
from app.retrieval.retrieval_service import hybrid_retrieve

REPORT_DIR = "generated_reports"

os.makedirs(REPORT_DIR, exist_ok=True)

def generate_docx_report(title: str, sections: list[str], prompt_context: str, tenant: str):
    doc = Document()
    now = datetime.now().strftime("%Y-%m-%d %H:%M")

    doc.add_heading(title or "Untitled Report", level=0)
    doc.add_paragraph(f"Tenant: {tenant} | Generated: {now}")
    doc.add_paragraph("Context: " + prompt_context)
    doc.add_page_break()

    for section in sections:
        doc.add_heading(section, level=1)

        # RAG-based answer generation
        rag_results = hybrid_retrieve(section, k=3)
        section_text = "\n".join([doc_chunk for doc_chunk, _ in rag_results]) or "No relevant data found."

        doc.add_paragraph(section_text)
        doc.add_page_break()

    # Save
    report_id = str(uuid.uuid4())
    file_path = f"{REPORT_DIR}/{report_id}.docx"
    doc.save(file_path)

    return report_id, file_path
